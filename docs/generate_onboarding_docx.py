"""Generate the staff onboarding Word document from the same source content
as docs/staff-onboarding.md.

Run::

    python3 generate_onboarding_docx.py

Writes ``docs/staff-onboarding.docx`` next to this script.
"""
from __future__ import annotations

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUT_PATH = os.path.join(os.path.dirname(__file__), "staff-onboarding.docx")

ENROLL_URL = "https://2mzabtr4o3vecbuembrosc7k2y0bdolo.lambda-url.us-east-1.on.aws/"
DASHBOARD_URL = "https://hl5bamdb5vdytk2p6mm527gyli0hxcrp.lambda-url.us-east-1.on.aws/"
ADMIN_EMAIL = "jma@nybrainspine.com"


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a background color to a table cell (python-docx has no helper)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink in the given paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2C6CDF")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _bullet(doc, text: str, *, level: int = 0):
    """Append a bulleted list paragraph using Word's built-in 'List Bullet'."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    return p


def build_document() -> None:
    doc = Document()

    # Tighter margins than Word's default 1" so the whole doc fits on one
    # printed page comfortably.
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Default body styling — slightly more compact than Word's defaults.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    # ── Title ─────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title_run = title.add_run("Getting Started with EMIS")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    title.paragraph_format.space_after = Pt(6)

    # Lead paragraph
    p = doc.add_paragraph()
    p.add_run(
        "EMIS (Email Ingestor & Scheduler) reads your Microsoft 365 inbox, "
        "calendar, and sent mail three times a week and emails you a "
        "structured agenda — what you owe people, what you're waiting on, "
        "what's coming up, and what to prepare for. You don't need to "
        "install anything; everything happens in the cloud."
    )

    p = doc.add_paragraph()
    p.add_run("You'll get three weekly emails:")

    _bullet(doc, "Monday 6:00 AM — Weekly agenda. The full picture for the week ahead.")
    _bullet(doc, "Wednesday 8:00 AM — Mid-week check-in. What's slipping, what needs attention by Friday.")
    _bullet(doc, "Friday 3:00 PM — End-of-week recap. What landed, what carries into next week.")

    p = doc.add_paragraph()
    p.add_run(
        "Plus an optional morning brief weekdays at 6:30 AM with prep notes "
        "for that day's meetings."
    )

    # ── Step 1: Enroll ────────────────────────────────────────────────
    h = doc.add_heading("Step 1: Enroll (one time, ~2 minutes)", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    p = doc.add_paragraph(style="List Number")
    p.add_run("Open this link in your browser: ")
    _add_hyperlink(p, ENROLL_URL, ENROLL_URL)
    doc.add_paragraph("Click Enroll with Microsoft 365 →", style="List Number")
    doc.add_paragraph("Sign in with your work email (yourname@nybrainspine.com)", style="List Number")
    doc.add_paragraph(
        "Review the permissions screen — EMIS asks to read your mail, calendar, "
        "tasks, and OneDrive files. It does not send mail on your behalf and "
        "cannot modify your data on its own.",
        style="List Number",
    )
    doc.add_paragraph('Click Accept. You\'ll see "You\'re enrolled!"', style="List Number")

    p = doc.add_paragraph()
    p.add_run("That's it. Your first agenda arrives on the next scheduled morning.")

    # ── Step 2: Read your first agenda ────────────────────────────────
    h = doc.add_heading("Step 2: Read your first agenda", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    p = doc.add_paragraph()
    p.add_run(
        "Your agenda email is a tidy summary. At the top there's a blue strip "
        "that says “View interactive dashboard →” — click it."
    )

    p = doc.add_paragraph()
    p.add_run("What you can do on the dashboard:").bold = True

    actions = [
        ("✓ done", "marks an item complete; it won't reappear next week"),
        ("1d · 1w · Mon", "snooze for 1 day, 1 week, or until next Monday"),
        ("✕", "drop the item permanently; EMIS won't surface it again"),
        ("📝 note", "add a free-text note (e.g., “told Sarah I'd send by Friday”) preserved in next week's agenda"),
        ("📌 pin", "pin an item to the top so it stays visible"),
        ("🖨 print view", "clean printable version for taking to a meeting"),
        ("Filter chips", "at the top of each section — show only clinical / business / admin / personal"),
    ]
    for label, desc in actions:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(label)
        r.bold = True
        p.add_run(" — " + desc)

    p = doc.add_paragraph()
    p.add_run("The dashboard URL is the same every time. Bookmark it: ")
    _add_hyperlink(p, DASHBOARD_URL, DASHBOARD_URL)
    p.add_run(" — you'll sign in with Microsoft 365 the first time and it remembers you for 7 days.")

    # ── Step 3: Adjust your settings ──────────────────────────────────
    h = doc.add_heading("Step 3: Adjust your settings", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    p = doc.add_paragraph()
    p.add_run("Click ")
    p.add_run("Settings").bold = True
    p.add_run(" in the top-right corner of the dashboard.")

    # Table of settings
    settings_rows = [
        ("Delivery channels", "Email is always on. SMS is currently disabled at the system level."),
        ("Schedules", "Uncheck a mode (Monday / Wednesday / Friday / Morning) to stop receiving that run."),
        ("Categories", "Pick which of clinical / business / admin / personal are relevant to your work. EMIS will surface those and de-emphasize the others."),
        ("Monthly spend cap", "Optional safety limit. If your scheduled runs exceed $X of AI cost in a calendar month, EMIS skips the rest of the month and emails you. Default: No cap."),
    ]
    table = doc.add_table(rows=len(settings_rows) + 1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.6)
    table.columns[1].width = Inches(5.4)

    hdr = table.rows[0].cells
    for c, label in zip(hdr, ("Setting", "What it does")):
        _set_cell_shading(c, "F0F4FA")
        run = c.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
    for i, (k, v) in enumerate(settings_rows, start=1):
        cells = table.rows[i].cells
        cells[0].paragraphs[0].add_run(k).bold = True
        cells[1].paragraphs[0].add_run(v)
        for c in cells:
            c.width = Inches(1.6) if c == cells[0] else Inches(5.4)

    p = doc.add_paragraph()
    r = p.add_run("Removing yourself: ")
    r.bold = True
    p.add_run(
        'The Danger zone at the bottom of Settings has a "Delete my account" '
        "button. Clicking it removes your enrollment and your stored agendas. "
        "Your Microsoft 365 inbox is untouched."
    )

    # ── Privacy & security ────────────────────────────────────────────
    h = doc.add_heading("Privacy & security", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    _bullet(doc, "EMIS runs entirely in AWS infrastructure covered by the practice's existing Business Associate Agreement.")
    _bullet(doc, "Your inbox content stays inside that infrastructure — it's never sent to outside vendors.")
    _bullet(doc, "The AI summary uses Claude (Anthropic) running on AWS Bedrock under HIPAA terms.")
    _bullet(doc, "Each enrolled user can only see their own dashboard. EMIS admins can see who's enrolled and total AI cost per person, but cannot read your agendas or your mail.")

    # ── Need help ─────────────────────────────────────────────────────
    h = doc.add_heading("Need help?", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    p = doc.add_paragraph()
    p.add_run("If something looks wrong or you don't get an expected email, contact John (")
    _add_hyperlink(p, f"mailto:{ADMIN_EMAIL}", ADMIN_EMAIL)
    p.add_run(").")

    p = doc.add_paragraph()
    p.add_run("If you ever forget the URLs:").italic = True

    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("Enroll: ")
    r.bold = True
    _add_hyperlink(p, ENROLL_URL, ENROLL_URL)

    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("Dashboard: ")
    r.bold = True
    _add_hyperlink(p, DASHBOARD_URL, DASHBOARD_URL)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build_document()
